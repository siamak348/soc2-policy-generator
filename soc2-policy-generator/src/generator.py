from pathlib import Path
from jinja2 import Environment, FileSystemLoader, select_autoescape
from docx import Document
from datetime import date

TEMPLATE_DIR = Path(__file__).parent / "policies"

def list_templates():
    return [p for p in TEMPLATE_DIR.glob("*.j2")]

def render_text(template_path, context):
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATE_DIR)),
        autoescape=select_autoescape(enabled_extensions=("j2",))
    )
    template = env.get_template(template_path.name)
    return template.render(**context)

def text_to_docx(text: str, outfile: Path):
    doc = Document()
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("**"))
            run.bold = True
        else:
            doc.add_paragraph(line)
    doc.save(str(outfile))

def generate_docs(output_dir: Path, selected_files, context):
    output_dir.mkdir(parents=True, exist_ok=True)
    for tpl in list_templates():
        if selected_files and tpl.name not in selected_files:
            continue
        text = render_text(tpl, context)
        out = output_dir / (tpl.stem + ".docx")
        text_to_docx(text, out)
    return output_dir